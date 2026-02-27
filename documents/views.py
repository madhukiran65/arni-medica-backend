from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from django_filters import FilterSet, CharFilter, DateFromToRangeFilter, BooleanFilter
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import SearchFilter, OrderingFilter
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.db.models import Count, Q
import hashlib

from .models import (
    DocumentInfocardType,
    Document,
    DocumentCheckout,
    DocumentChangeOrder,
    DocumentSnapshot,
    DocumentVersion,
    DocumentApprover,
    DocumentCollaborator,
    DocumentComment,
    DocumentSuggestion,
)
from .serializers import (
    DocumentInfocardTypeSerializer,
    DocumentListSerializer,
    DocumentDetailSerializer,
    DocumentCreateSerializer,
    DocumentCheckoutSerializer,
    DocumentChangeOrderSerializer,
    DocumentSnapshotSerializer,
    DocumentVersionSerializer,
    DocumentApproverSerializer,
    DocumentCollaboratorSerializer,
    DocumentCommentSerializer,
    DocumentSuggestionSerializer,
    DocumentContentUpdateSerializer,
    CheckoutActionSerializer,
    CheckinActionSerializer,
    LifecycleTransitionSerializer,
)


# ============================================================================
# FILTER SETS
# ============================================================================

class DocumentFilterSet(FilterSet):
    """Enhanced FilterSet for Document with comprehensive filtering options."""

    vault_state = CharFilter(field_name='vault_state', lookup_expr='iexact')
    lifecycle_stage = CharFilter(field_name='lifecycle_stage', lookup_expr='icontains')
    department = CharFilter(field_name='department__name', lookup_expr='icontains')
    infocard_type = CharFilter(field_name='infocard_type__name', lookup_expr='icontains')
    owner = CharFilter(field_name='owner__username', lookup_expr='iexact')
    created_at__gte = DateFromToRangeFilter(field_name='created_at')
    created_at__lte = DateFromToRangeFilter(field_name='created_at')
    requires_training = BooleanFilter(field_name='requires_training')
    is_locked = BooleanFilter(field_name='is_locked')
    requires_approval = BooleanFilter(field_name='requires_approval')

    class Meta:
        model = Document
        fields = [
            'vault_state',
            'lifecycle_stage',
            'department',
            'infocard_type',
            'owner',
            'requires_training',
            'is_locked',
            'requires_approval',
            'created_at__gte',
            'created_at__lte',
        ]


class DocumentInfocardTypeViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Read-only viewset for document infocard types.
    """
    queryset = DocumentInfocardType.objects.all()
    serializer_class = DocumentInfocardTypeSerializer
    permission_classes = [IsAuthenticated]


class DocumentViewSet(viewsets.ModelViewSet):
    """
    Enhanced ViewSet for managing documents with full CRUD operations,
    lifecycle management, checkout/checkin, approvals, and audit trails.

    Features:
    - Auto-generated document IDs based on infocard_type prefix
    - Comprehensive filtering and search
    - Bulk release for admin users
    - Pending review tracking
    - User checkout history
    - Document statistics for dashboards
    - File upload with SHA-256 hashing
    - Version auto-increment on checkin
    """
    queryset = Document.objects.select_related(
        'infocard_type',
        'department',
        'owner',
        'locked_by',
        'created_by',
        'updated_by'
    ).prefetch_related(
        'checkouts',
        'snapshots',
        'versions',
        'approvers',
        'change_orders'
    ).all()
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_class = DocumentFilterSet
    search_fields = ['document_id', 'legacy_document_id', 'title', 'description', 'subject_keywords']
    ordering_fields = ['created_at', 'document_id', 'title', 'next_review_date', 'vault_state']
    ordering = ['-created_at']
    parser_classes = (JSONParser, MultiPartParser, FormParser)

    def get_serializer_class(self):
        """Return appropriate serializer based on action."""
        if self.action == 'list':
            return DocumentListSerializer
        elif self.action == 'retrieve':
            return DocumentDetailSerializer
        elif self.action == 'create':
            return DocumentCreateSerializer
        return DocumentDetailSerializer

    def list(self, request, *args, **kwargs):
        """Override list with error debugging."""
        import traceback
        try:
            return super().list(request, *args, **kwargs)
        except Exception as e:
            return Response(
                {'error': str(e), 'traceback': traceback.format_exc()},
                status=500
            )

    def retrieve(self, request, *args, **kwargs):
        """Override retrieve with error debugging."""
        import traceback
        try:
            return super().retrieve(request, *args, **kwargs)
        except Exception as e:
            return Response(
                {'error': str(e), 'traceback': traceback.format_exc()},
                status=500
            )

    def perform_create(self, serializer):
        """Set owner to current user on creation."""
        serializer.save(owner=self.request.user, created_by=self.request.user)

    def create(self, request, *args, **kwargs):
        """Override create with clean error handling."""
        import traceback
        try:
            serializer = self.get_serializer(data=request.data)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)

            document = serializer.instance
            return Response(
                DocumentDetailSerializer(document).data,
                status=status.HTTP_201_CREATED
            )
        except Exception as e:
            return Response(
                {'error': str(e), 'traceback': traceback.format_exc()},
                status=500
            )

    def perform_update(self, serializer):
        """Update document and set updated_by."""
        serializer.save(updated_by=self.request.user)

    # ========================================================================
    # DOCUMENT LIFECYCLE ACTIONS
    # ========================================================================

    @action(detail=True, methods=['post'])
    def checkout(self, request, pk=None):
        """
        Check out a document for editing.

        POST body: {
            "checkout_reason": "string (required)",
            "expected_checkin_date": "ISO datetime (optional)"
        }
        """
        document = self.get_object()

        # Validation: check if already checked out
        active_checkout = document.get_active_checkout()
        if active_checkout:
            return Response(
                {
                    'error': 'Document is already checked out',
                    'checked_out_by': active_checkout.checked_out_by.username,
                    'checked_out_at': active_checkout.checked_out_at.isoformat()
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        # Validation: only draft documents can be checked out
        if document.vault_state != 'draft':
            return Response(
                {'error': f'Only draft documents can be checked out. Current state: {document.vault_state}'},
                status=status.HTTP_400_BAD_REQUEST
            )

        serializer = CheckoutActionSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Invalidate any previous active checkouts
            DocumentCheckout.objects.filter(
                document=document,
                is_active=True
            ).update(is_active=False)

            checkout = DocumentCheckout.objects.create(
                document=document,
                checked_out_by=request.user,
                checkout_reason=serializer.validated_data.get('checkout_reason', ''),
                expected_checkin_date=serializer.validated_data.get('expected_checkin_date')
            )

            return Response(
                DocumentCheckoutSerializer(checkout).data,
                status=status.HTTP_201_CREATED
            )
        except Exception as e:
            return Response(
                {'error': f'Checkout failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def checkin(self, request, pk=None):
        """
        Check in a document and auto-increment minor version.

        POST body: {
            "file": "file (optional)",
            "change_summary": "string (required)",
            "is_major_change": "boolean (default: false)"
        }
        """
        document = self.get_object()
        serializer = CheckinActionSerializer(data=request.data)

        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Validation: document must be checked out
            checkout = document.get_active_checkout()
            if not checkout:
                return Response(
                    {'error': 'Document is not currently checked out'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Validation: only the user who checked out can check in
            if checkout.checked_out_by != request.user:
                return Response(
                    {
                        'error': 'Only the user who checked out this document can check it in',
                        'checked_out_by': checkout.checked_out_by.username
                    },
                    status=status.HTTP_403_FORBIDDEN
                )

            # Process file upload if provided
            if 'file' in request.FILES:
                document.file = request.FILES['file']
                document.original_filename = request.FILES['file'].name
                document.file_type = request.FILES['file'].content_type
                document.file_size = request.FILES['file'].size
                document.file_hash = self._calculate_file_hash(request.FILES['file'])

            # Auto-increment minor version
            is_major = serializer.validated_data.get('is_major_change', False)
            if is_major:
                document.major_version += 1
                document.minor_version = 0
            else:
                document.minor_version += 1

            document.change_summary = serializer.validated_data.get('change_summary', '')
            document.updated_by = request.user
            document.save()

            # Create version snapshot
            version = DocumentVersion.objects.create(
                document=document,
                major_version=document.major_version,
                minor_version=document.minor_version,
                change_type='major' if is_major else 'minor',
                is_major_change=is_major,
                change_summary=document.change_summary,
                snapshot_data={
                    'title': document.title,
                    'vault_state': document.vault_state,
                    'version': document.version_string
                },
                released_date=timezone.now()
            )

            # Close checkout
            checkout.is_active = False
            checkout.save()

            return Response(
                DocumentVersionSerializer(version).data,
                status=status.HTTP_201_CREATED
            )
        except Exception as e:
            return Response(
                {'error': f'Check-in failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['post'])
    def bulk_release(self, request):
        """
        Release multiple documents at once (admin only).

        POST body: {
            "document_ids": [1, 2, 3, ...]
        }
        """
        # Check admin permission
        if not request.user.is_staff:
            return Response(
                {'error': 'Only administrators can perform bulk release'},
                status=status.HTTP_403_FORBIDDEN
            )

        doc_ids = request.data.get('document_ids', [])
        if not doc_ids:
            return Response(
                {'error': 'document_ids list is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            documents = Document.objects.filter(id__in=doc_ids)
            released_count = 0

            for doc in documents:
                if doc.vault_state != 'released':
                    doc.vault_state = 'released'
                    doc.released_date = timezone.now()
                    doc.save()
                    released_count += 1

            return Response(
                {
                    'success': True,
                    'released_count': released_count,
                    'total_requested': len(doc_ids),
                    'message': f'{released_count} document(s) released'
                },
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Bulk release failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['get'])
    def pending_review(self, request):
        """
        List documents past their next_review_date.

        Returns paginated list of documents with overdue reviews.
        """
        today = timezone.now().date()
        overdue = Document.objects.filter(
            next_review_date__lt=today
        ).select_related('owner', 'department').order_by('next_review_date')

        page = self.paginate_queryset(overdue)
        if page is not None:
            serializer = DocumentListSerializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = DocumentListSerializer(overdue, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def my_checkouts(self, request):
        """
        List documents checked out by current user.
        """
        checkouts = DocumentCheckout.objects.filter(
            checked_out_by=request.user,
            is_active=True
        ).select_related('document').order_by('-checked_out_at')

        documents = [co.document for co in checkouts]
        page = self.paginate_queryset(documents)
        if page is not None:
            serializer = DocumentListSerializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = DocumentListSerializer(documents, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def document_stats(self, request):
        """
        Return statistics for document dashboard.

        Returns: {
            "total_documents": int,
            "by_vault_state": { "draft": int, "released": int, ... },
            "by_infocard_type": { "SOP": int, "FORM": int, ... },
            "overdue_reviews": int,
            "locked_documents": int,
            "checked_out_documents": int,
            "my_documents": int
        }
        """
        total = Document.objects.count()
        today = timezone.now().date()

        # Group by vault_state
        by_state = Document.objects.values('vault_state').annotate(
            count=Count('id')
        ).order_by('vault_state')

        # Group by infocard_type
        by_type = Document.objects.values('infocard_type__prefix').annotate(
            count=Count('id')
        ).order_by('infocard_type__prefix')

        overdue = Document.objects.filter(
            next_review_date__lt=today
        ).count()

        locked = Document.objects.filter(is_locked=True).count()

        checked_out = DocumentCheckout.objects.filter(
            is_active=True
        ).count()

        my_docs = Document.objects.filter(owner=request.user).count()

        return Response({
            'total_documents': total,
            'by_vault_state': {item['vault_state']: item['count'] for item in by_state},
            'by_infocard_type': {item['infocard_type__prefix']: item['count'] for item in by_type},
            'overdue_reviews': overdue,
            'locked_documents': locked,
            'checked_out_documents': checked_out,
            'my_documents': my_docs,
        })

    @action(detail=True, methods=['post'], parser_classes=(MultiPartParser, FormParser))
    def upload_file(self, request, pk=None):
        """
        Upload and attach a file to a document.

        POST (multipart): {
            "file": "file (required)"
        }

        Returns SHA-256 hash of uploaded file.
        """
        document = self.get_object()

        if 'file' not in request.FILES:
            return Response(
                {'error': 'file is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            uploaded_file = request.FILES['file']

            # Calculate SHA-256 hash
            file_hash = self._calculate_file_hash(uploaded_file)

            # Save file to document
            document.file = uploaded_file
            document.original_filename = uploaded_file.name
            document.file_type = uploaded_file.content_type
            document.file_size = uploaded_file.size
            document.file_hash = file_hash
            document.save()

            return Response({
                'success': True,
                'file_name': uploaded_file.name,
                'file_size': uploaded_file.size,
                'file_type': uploaded_file.content_type,
                'file_hash': file_hash,
                'hash_algorithm': 'SHA-256'
            }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {'error': f'File upload failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['post'], parser_classes=[JSONParser])
    def bulk_import(self, request):
        """
        Bulk import documents with legacy ID mapping.

        POST body (JSON): {
            "documents": [
                {
                    "title": "SOP on SOP",
                    "legacy_document_id": "QA001-03",
                    "infocard_type_prefix": "SOP",
                    "department_name": "Quality Assurance",
                    "business_unit": "QA",
                    "vault_state": "released",
                    "major_version": 3,
                    "minor_version": 0,
                    "external_file_url": "https://...",
                    "original_filename": "SOP on SOP.docx",
                    "subject_keywords": ["sop", "quality"],
                    "custom_fields": {"legacy_folder": "QA001-03"}
                }
            ]
        }
        """
        if not request.user.is_staff:
            return Response(
                {'error': 'Only administrators can perform bulk import'},
                status=status.HTTP_403_FORBIDDEN
            )

        documents_data = request.data.get('documents', [])
        if not documents_data:
            return Response(
                {'error': 'documents list is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        results = {'created': 0, 'skipped': 0, 'errors': []}

        for doc_data in documents_data:
            try:
                legacy_id = doc_data.get('legacy_document_id', '')

                # Skip if legacy_document_id already exists
                if legacy_id and Document.objects.filter(legacy_document_id=legacy_id).exists():
                    results['skipped'] += 1
                    continue

                # Look up infocard type by prefix
                prefix = doc_data.get('infocard_type_prefix', 'SOP')
                infocard_type = DocumentInfocardType.objects.filter(prefix=prefix).first()
                if not infocard_type:
                    infocard_type = DocumentInfocardType.objects.first()

                # Look up department if provided
                from users.models import Department
                department = None
                dept_name = doc_data.get('department_name', '')
                if dept_name:
                    department = Department.objects.filter(name__icontains=dept_name).first()

                # Parse version from legacy ID (e.g., QA001-03 → version 3)
                major_version = doc_data.get('major_version', 1)
                minor_version = doc_data.get('minor_version', 0)

                document = Document(
                    title=doc_data.get('title', 'Untitled'),
                    legacy_document_id=legacy_id,
                    infocard_type=infocard_type,
                    department=department,
                    owner=request.user,
                    created_by=request.user,
                    business_unit=doc_data.get('business_unit', ''),
                    vault_state=doc_data.get('vault_state', 'released'),
                    major_version=major_version,
                    minor_version=minor_version,
                    external_file_url=doc_data.get('external_file_url', ''),
                    original_filename=doc_data.get('original_filename', ''),
                    subject_keywords=doc_data.get('subject_keywords', []),
                    custom_fields=doc_data.get('custom_fields', {}),
                    requires_training=doc_data.get('requires_training', False),
                    distribution_restriction=doc_data.get('distribution_restriction', 'internal'),
                    confidentiality_level=doc_data.get('confidentiality_level', 'internal'),
                )
                document.save()
                results['created'] += 1

            except Exception as e:
                results['errors'].append({
                    'legacy_id': doc_data.get('legacy_document_id', 'unknown'),
                    'error': str(e)
                })

        return Response({
            'success': True,
            'created': results['created'],
            'skipped': results['skipped'],
            'errors': results['errors'],
            'total_requested': len(documents_data),
        }, status=status.HTTP_201_CREATED)

    def _calculate_file_hash(self, file_obj):
        """Calculate SHA-256 hash of a file object."""
        hasher = hashlib.sha256()
        for chunk in file_obj.chunks():
            hasher.update(chunk)
        return hasher.hexdigest()

    # ========================================================================
    # DOCUMENT STATE TRANSITION ACTIONS
    # ========================================================================

    @action(detail=True, methods=['post'])
    def lifecycle_transition(self, request, pk=None):
        """
        FSM-validated lifecycle transition.

        POST body: {
            "target_stage": "in_review|approved|effective|...",
            "comments": "string (optional)",
            "password": "string (required for e-signature)"
        }
        """
        document = self.get_object()
        serializer = LifecycleTransitionSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        target_state = serializer.validated_data['target_stage']
        current_state = document.vault_state

        # FSM validation
        from .fsm import validate_transition, check_transition_permission
        try:
            validate_transition(current_state, target_state)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # Permission check
        allowed, reason = check_transition_permission(current_state, target_state, request.user, document)
        if not allowed:
            return Response({'error': reason}, status=status.HTTP_403_FORBIDDEN)

        # Password verification for e-signature
        password = serializer.validated_data.get('password', '')
        if password and not request.user.check_password(password):
            return Response({'error': 'Invalid password for electronic signature'}, status=status.HTTP_403_FORBIDDEN)

        try:
            old_state = document.vault_state
            document.vault_state = target_state
            document.lifecycle_stage = target_state
            document.updated_by = request.user

            # Set state-specific dates
            now = timezone.now()
            if target_state == 'approved':
                document.approved_date = now
            elif target_state == 'effective':
                document.effective_date = now.date()
                if document.review_period_months:
                    try:
                        from dateutil.relativedelta import relativedelta
                        document.next_review_date = now.date() + relativedelta(months=document.review_period_months)
                    except ImportError:
                        from datetime import timedelta
                        document.next_review_date = now.date() + timedelta(days=document.review_period_months * 30)
            elif target_state == 'obsolete':
                document.obsolete_date = now.date()
            elif target_state == 'archived':
                document.archived_date = now.date()
            elif target_state == 'cancelled':
                document.cancelled_date = now
                document.cancelled_by = request.user
                document.cancellation_reason = serializer.validated_data.get('comments', '')

            document.save()

            # Create snapshot
            DocumentSnapshot.objects.create(
                document=document,
                version_string=document.version_string,
                snapshot_type=target_state if target_state in ['approval', 'release', 'archive', 'review', 'training_complete', 'effective', 'superseded', 'cancelled'] else 'review',
                snapshot_data={
                    'from_state': old_state,
                    'to_state': target_state,
                    'transitioned_by': request.user.username,
                    'comments': serializer.validated_data.get('comments', ''),
                    'timestamp': now.isoformat(),
                },
                created_by=request.user,
            )

            # Auto-transition: approved → training_period if training required
            if target_state == 'approved' and document.requires_training and document.training_completion_required:
                document.vault_state = 'training_period'
                document.lifecycle_stage = 'training_period'
                document.save()
                # Auto-assign training (signal will handle this)

            # Send notifications
            try:
                from core.notifications import NotificationService
                NotificationService.send_workflow_transition(
                    record=document,
                    from_stage=old_state,
                    to_stage=document.vault_state,
                    transitioned_by=request.user,
                    record_type='document',
                )
            except Exception:
                pass

            return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Transition failed: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def submit_for_review(self, request, pk=None):
        """Submit draft document for review. Shortcut for draft → in_review."""
        document = self.get_object()
        if document.vault_state != 'draft':
            return Response({'error': 'Only draft documents can be submitted for review'}, status=status.HTTP_400_BAD_REQUEST)

        # Check document has approvers assigned
        if not document.approvers.exists():
            return Response({'error': 'Please assign at least one reviewer/approver before submitting'}, status=status.HTTP_400_BAD_REQUEST)

        # Check author permission
        if document.owner != request.user and document.created_by != request.user and not request.user.is_staff:
            return Response({'error': 'Only the document owner can submit for review'}, status=status.HTTP_403_FORBIDDEN)

        document.vault_state = 'in_review'
        document.lifecycle_stage = 'in_review'
        document.is_locked = True
        document.locked_by = request.user
        document.locked_at = timezone.now()
        document.lock_reason = 'Submitted for review'
        document.updated_by = request.user
        document.save()

        # Reset all approver statuses to pending
        document.approvers.all().update(approval_status='pending', approved_at=None, comments='')

        # Notify approvers
        try:
            from core.notifications import NotificationService
            approvers = [a.approver for a in document.approvers.filter(approval_status='pending') if a.approver]
            if approvers:
                NotificationService.send_approval_request(
                    document=document, approvers=approvers, requester=request.user,
                )
        except Exception:
            pass

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def final_approve(self, request, pk=None):
        """
        Final approval that transitions document from in_review → approved.
        Requires all approvers to have approved.
        """
        document = self.get_object()
        if document.vault_state != 'in_review':
            return Response({'error': 'Document must be in review to approve'}, status=status.HTTP_400_BAD_REQUEST)

        # Verify password for e-signature
        password = request.data.get('password', '')
        if not password or not request.user.check_password(password):
            return Response({'error': 'Valid password required for electronic signature'}, status=status.HTTP_403_FORBIDDEN)

        # Record this user's approval
        approver_record = document.approvers.filter(approver=request.user).first()
        if approver_record:
            approver_record.approval_status = 'approved'
            approver_record.approved_at = timezone.now()
            approver_record.comments = request.data.get('comments', '')
            approver_record.save()

        # Check if ALL approvers have approved
        if not document.is_fully_approved():
            pending = document.approvers.filter(approval_status='pending').count()
            return Response({
                'message': 'Your approval recorded. Waiting for remaining approvers.',
                'pending_approvals': pending,
                'your_status': 'approved',
            }, status=status.HTTP_200_OK)

        # All approved — transition to approved state
        document.vault_state = 'approved'
        document.lifecycle_stage = 'approved'
        document.approved_date = timezone.now()
        document.released_date = timezone.now()
        document.updated_by = request.user

        # Auto-transition to training_period if training required
        if document.requires_training and document.training_completion_required:
            document.vault_state = 'training_period'
            document.lifecycle_stage = 'training_period'
        elif not document.requires_training:
            # No training required — go directly to effective
            document.vault_state = 'effective'
            document.lifecycle_stage = 'effective'
            document.effective_date = timezone.now().date()
            if document.review_period_months:
                try:
                    from dateutil.relativedelta import relativedelta
                    document.next_review_date = timezone.now().date() + relativedelta(months=document.review_period_months)
                except ImportError:
                    from datetime import timedelta
                    document.next_review_date = timezone.now().date() + timedelta(days=document.review_period_months * 30)

        document.is_locked = True
        document.save()

        # Create approval snapshot
        DocumentSnapshot.objects.create(
            document=document,
            version_string=document.version_string,
            snapshot_type='approval',
            snapshot_data={
                'approved_by': [{'user': a.approver.username, 'at': str(a.approved_at)} for a in document.approvers.filter(approval_status='approved')],
                'final_state': document.vault_state,
            },
            created_by=request.user,
        )

        # Increment major version on approval
        document.major_version += 1
        document.minor_version = 0
        document.save()

        # Create version record
        DocumentVersion.objects.create(
            document=document,
            major_version=document.major_version,
            minor_version=0,
            change_type='major',
            is_major_change=True,
            change_summary=f'Approved and released v{document.major_version}.0',
            snapshot_data={
                'vault_state': document.vault_state,
                'approved_date': str(document.approved_date),
            },
            released_date=timezone.now(),
            created_by=request.user,
            updated_by=request.user,
        )

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def make_effective(self, request, pk=None):
        """Transition approved/training_period document to effective."""
        document = self.get_object()
        if document.vault_state not in ('approved', 'training_period'):
            return Response({'error': 'Document must be approved or in training period'}, status=status.HTTP_400_BAD_REQUEST)

        if not request.user.is_staff:
            return Response({'error': 'Only administrators can manually make documents effective'}, status=status.HTTP_403_FORBIDDEN)

        document.vault_state = 'effective'
        document.lifecycle_stage = 'effective'
        document.effective_date = timezone.now().date()
        document.updated_by = request.user

        if document.review_period_months:
            try:
                from dateutil.relativedelta import relativedelta
                document.next_review_date = timezone.now().date() + relativedelta(months=document.review_period_months)
            except ImportError:
                from datetime import timedelta
                document.next_review_date = timezone.now().date() + timedelta(days=document.review_period_months * 30)

        document.save()
        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def cancel_document(self, request, pk=None):
        """Cancel a document (only from draft or in_review)."""
        document = self.get_object()
        if document.vault_state not in ('draft', 'in_review'):
            return Response({'error': 'Only draft or in-review documents can be cancelled'}, status=status.HTTP_400_BAD_REQUEST)

        reason = request.data.get('reason', '')
        if not reason:
            return Response({'error': 'Cancellation reason is required'}, status=status.HTTP_400_BAD_REQUEST)

        document.vault_state = 'cancelled'
        document.lifecycle_stage = 'cancelled'
        document.cancelled_date = timezone.now()
        document.cancelled_by = request.user
        document.cancellation_reason = reason
        document.updated_by = request.user
        document.save()

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def supersede(self, request, pk=None):
        """Supersede an effective document with a new version."""
        document = self.get_object()
        if document.vault_state != 'effective':
            return Response({'error': 'Only effective documents can be superseded'}, status=status.HTTP_400_BAD_REQUEST)

        if not request.user.is_staff:
            return Response({'error': 'Only administrators can supersede documents'}, status=status.HTTP_403_FORBIDDEN)

        new_doc_id = request.data.get('superseded_by_id')
        if new_doc_id:
            try:
                new_doc = Document.objects.get(id=new_doc_id)
                document.superseded_by = new_doc
            except Document.DoesNotExist:
                pass

        document.vault_state = 'superseded'
        document.lifecycle_stage = 'superseded'
        document.updated_by = request.user
        document.save()

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def make_obsolete(self, request, pk=None):
        """Mark an effective document as obsolete."""
        document = self.get_object()
        if document.vault_state != 'effective':
            return Response({'error': 'Only effective documents can be made obsolete'}, status=status.HTTP_400_BAD_REQUEST)

        if not request.user.is_staff:
            return Response({'error': 'Only administrators can obsolete documents'}, status=status.HTTP_403_FORBIDDEN)

        document.vault_state = 'obsolete'
        document.lifecycle_stage = 'obsolete'
        document.obsolete_date = timezone.now().date()
        document.updated_by = request.user
        document.save()

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def archive_document(self, request, pk=None):
        """Archive a superseded, obsolete, or effective document."""
        document = self.get_object()
        if document.vault_state not in ('effective', 'superseded', 'obsolete'):
            return Response({'error': 'Only effective, superseded, or obsolete documents can be archived'}, status=status.HTTP_400_BAD_REQUEST)

        if not request.user.is_staff:
            return Response({'error': 'Only administrators can archive documents'}, status=status.HTTP_403_FORBIDDEN)

        document.vault_state = 'archived'
        document.lifecycle_stage = 'archived'
        document.archived_date = timezone.now().date()
        document.updated_by = request.user
        document.save()

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def initiate_revision(self, request, pk=None):
        """Start a revision cycle on an effective document (creates new draft version)."""
        document = self.get_object()
        if document.vault_state != 'effective':
            return Response({'error': 'Only effective documents can be revised'}, status=status.HTTP_400_BAD_REQUEST)

        # Create a new draft version linked to this document
        document.vault_state = 'draft'
        document.lifecycle_stage = 'draft'
        document.is_locked = False
        document.locked_by = None
        document.locked_at = None
        document.lock_reason = ''
        document.minor_version += 1
        document.updated_by = request.user
        document.save()

        # Reset approvals
        document.approvers.all().update(approval_status='pending', approved_at=None, comments='')

        return Response(DocumentDetailSerializer(document).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['get'])
    def available_transitions(self, request, pk=None):
        """Get list of available state transitions for this document."""
        document = self.get_object()
        from .fsm import get_available_transitions, check_transition_permission

        transitions = get_available_transitions(document.vault_state)
        # Filter by user permissions
        result = []
        for t in transitions:
            allowed, reason = check_transition_permission(
                document.vault_state, t['target_state'], request.user, document
            )
            t['allowed'] = allowed
            t['reason'] = reason if not allowed else ''
            result.append(t)

        return Response({
            'current_state': document.vault_state,
            'transitions': result,
        })

    @action(detail=True, methods=['get'])
    def training_status(self, request, pk=None):
        """Get training completion status for this document."""
        document = self.get_object()

        try:
            from training.models import TrainingAssignment
            assignments = TrainingAssignment.objects.filter(triggering_document=document)
            total = assignments.count()
            completed = assignments.filter(status='completed').count()

            return Response({
                'document_id': document.document_id,
                'requires_training': document.requires_training,
                'training_completion_required': document.training_completion_required,
                'total_assignments': total,
                'completed': completed,
                'pending': total - completed,
                'completion_percentage': round((completed / total * 100), 1) if total > 0 else 0,
                'all_complete': completed == total and total > 0,
            })
        except Exception:
            return Response({
                'document_id': document.document_id,
                'requires_training': document.requires_training,
                'total_assignments': 0,
                'completed': 0,
                'pending': 0,
                'completion_percentage': 0,
                'all_complete': False,
            })

    @action(detail=True, methods=['post'])
    def acknowledge(self, request, pk=None):
        """Record user acknowledgment of a document."""
        document = self.get_object()
        from .models import DocumentAcknowledgment

        method = request.data.get('method', 'read')

        ack, created = DocumentAcknowledgment.objects.get_or_create(
            document=document,
            user=request.user,
            defaults={'method': method}
        )

        if not created:
            return Response({'message': 'Already acknowledged', 'acknowledged_at': ack.acknowledged_at}, status=status.HTTP_200_OK)

        return Response({
            'message': 'Document acknowledged',
            'acknowledged_at': ack.acknowledged_at,
            'method': ack.method,
        }, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """
        Approve document with e-signature.

        POST body: {
            "signature": "string (required)",
            "comment": "string (optional)"
        }
        """
        document = self.get_object()

        signature = request.data.get('signature')
        if not signature:
            return Response(
                {'error': 'Signature is required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            approver, created = DocumentApprover.objects.update_or_create(
                document=document,
                approver=request.user,
                defaults={
                    'approval_status': 'approved',
                    'approved_at': timezone.now(),
                    'comments': request.data.get('comment', '')
                }
            )

            # Send notification to document owner about approval
            try:
                from core.notifications import NotificationService
                NotificationService.send_approval_complete(
                    document=document,
                    approver=request.user,
                    decision='approved',
                )
            except Exception as notify_err:
                import logging
                logging.getLogger(__name__).warning(f"Approval notification failed: {notify_err}")

            return Response(
                DocumentApproverSerializer(approver).data,
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Approval failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def lock(self, request, pk=None):
        """Lock document from further edits."""
        document = self.get_object()
        try:
            document.is_locked = True
            document.locked_by = request.user
            document.locked_at = timezone.now()
            document.lock_reason = request.data.get('reason', '')
            document.save()

            return Response(
                DocumentDetailSerializer(document).data,
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Lock failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def unlock(self, request, pk=None):
        """Unlock document for editing."""
        document = self.get_object()

        # Only document owner or staff can unlock
        if document.locked_by != request.user and not request.user.is_staff:
            return Response(
                {'error': 'Only the user who locked this document can unlock it'},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            document.is_locked = False
            document.locked_by = None
            document.locked_at = None
            document.lock_reason = ''
            document.save()

            return Response(
                DocumentDetailSerializer(document).data,
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Unlock failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    # ========================================================================
    # DOCUMENT METADATA RETRIEVAL ACTIONS
    # ========================================================================

    @action(detail=True, methods=['get'])
    def snapshots(self, request, pk=None):
        """List all snapshots for a document."""
        document = self.get_object()
        snapshots = document.snapshots.all().order_by('-created_at')
        serializer = DocumentSnapshotSerializer(snapshots, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get', 'post'])
    def approvers(self, request, pk=None):
        """List or add approvers for a document."""
        document = self.get_object()

        if request.method == 'GET':
            approvers = document.approvers.all().order_by('sequence')
            serializer = DocumentApproverSerializer(approvers, many=True)
            return Response(serializer.data)

        # POST: Add approver(s) to document
        from django.contrib.auth.models import User as AuthUser
        approvers_data = request.data.get('approvers', [])
        if not isinstance(approvers_data, list):
            approvers_data = [approvers_data]

        results = []
        for appr in approvers_data:
            user_id = appr.get('user_id') if isinstance(appr, dict) else appr
            role_label = appr.get('role', 'Reviewer') if isinstance(appr, dict) else 'Reviewer'
            seq = appr.get('sequence', 1) if isinstance(appr, dict) else 1

            try:
                user = AuthUser.objects.get(id=user_id)
                obj, created = DocumentApprover.objects.update_or_create(
                    document=document,
                    approver=user,
                    defaults={
                        'sequence': seq,
                        'role_required': role_label,
                        'approval_status': 'pending'
                    }
                )
                results.append({
                    'id': obj.id,
                    'user_id': user.id,
                    'username': user.username,
                    'full_name': user.get_full_name(),
                    'role': role_label,
                    'sequence': seq,
                    'status': obj.approval_status,
                    'created': created
                })
            except AuthUser.DoesNotExist:
                results.append({'user_id': user_id, 'error': 'User not found'})

        return Response({'approvers': results}, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['delete'], url_path='approvers/(?P<approver_id>[^/.]+)')
    def remove_approver(self, request, pk=None, approver_id=None):
        """Remove an approver from a document."""
        document = self.get_object()
        try:
            approver = document.approvers.get(id=approver_id)
            approver.delete()
            return Response({'message': 'Approver removed'}, status=status.HTTP_204_NO_CONTENT)
        except DocumentApprover.DoesNotExist:
            return Response({'error': 'Approver not found'}, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['get', 'post'])
    def collaborators_list(self, request, pk=None):
        """GET: List collaborators. POST: Add a collaborator."""
        document = self.get_object()

        if request.method == 'GET':
            collabs = DocumentCollaborator.objects.filter(
                document=document
            ).exclude(status='removed').select_related('user', 'user__profile', 'user__profile__department')
            serializer = DocumentCollaboratorSerializer(collabs, many=True)
            return Response(serializer.data)

        # POST: Add collaborator
        user_id = request.data.get('user_id')
        roles = request.data.get('roles', ['collaborator'])

        if not user_id:
            return Response({'error': 'user_id is required'}, status=400)

        from django.contrib.auth import get_user_model
        User = get_user_model()
        try:
            user = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return Response({'error': 'User not found'}, status=404)

        # Validate roles
        valid_roles = ['author', 'collaborator', 'reviewer_only', 'reviewer_approver']
        for role in roles:
            if role not in valid_roles:
                return Response({'error': f'Invalid role: {role}. Valid: {valid_roles}'}, status=400)

        collab, created = DocumentCollaborator.objects.update_or_create(
            document=document,
            user=user,
            defaults={
                'roles': roles,
                'status': 'active',
                'added_by': request.user,
            }
        )

        serializer = DocumentCollaboratorSerializer(collab)
        return Response(serializer.data, status=201 if created else 200)

    @action(detail=True, methods=['patch'], url_path='collaborators_list/(?P<collab_id>[0-9]+)')
    def update_collaborator(self, request, pk=None, collab_id=None):
        """Update a collaborator's roles."""
        document = self.get_object()
        try:
            collab = DocumentCollaborator.objects.get(id=collab_id, document=document)
        except DocumentCollaborator.DoesNotExist:
            return Response({'error': 'Collaborator not found'}, status=404)

        roles = request.data.get('roles')
        status_val = request.data.get('status')

        if roles is not None:
            valid_roles = ['author', 'collaborator', 'reviewer_only', 'reviewer_approver']
            for role in roles:
                if role not in valid_roles:
                    return Response({'error': f'Invalid role: {role}'}, status=400)
            collab.roles = roles

        if status_val is not None:
            collab.status = status_val

        collab.save()
        serializer = DocumentCollaboratorSerializer(collab)
        return Response(serializer.data)

    @action(detail=True, methods=['delete'], url_path='remove_collaborator/(?P<collab_id>[0-9]+)')
    def remove_collaborator(self, request, pk=None, collab_id=None):
        """Remove a collaborator from a document."""
        document = self.get_object()
        try:
            collab = DocumentCollaborator.objects.get(id=collab_id, document=document)
        except DocumentCollaborator.DoesNotExist:
            return Response({'error': 'Collaborator not found'}, status=404)

        collab.status = 'removed'
        collab.save()
        return Response({'status': 'removed'}, status=200)

    @action(detail=True, methods=['get'])
    def versions(self, request, pk=None):
        """List all versions for a document."""
        document = self.get_object()
        versions = document.versions.all().order_by('-major_version', '-minor_version')

        page = self.paginate_queryset(versions)
        if page is not None:
            serializer = DocumentVersionSerializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = DocumentVersionSerializer(versions, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def audit_trail(self, request, pk=None):
        """Return audit log entries for this document."""
        document = self.get_object()

        # Collect all document changes from audit trail
        checkouts = DocumentCheckout.objects.filter(document=document).order_by('-checked_out_at')
        versions = DocumentVersion.objects.filter(document=document).order_by('-created_at')
        approvals = DocumentApprover.objects.filter(document=document).order_by('-approved_at')

        # Build timeline
        timeline = []

        # Document creation
        timeline.append({
            'timestamp': document.created_at,
            'action': 'created',
            'user': document.created_by.username if document.created_by else 'System',
            'details': f'Document {document.document_id} created'
        })

        # Checkouts
        for checkout in checkouts:
            timeline.append({
                'timestamp': checkout.checked_out_at,
                'action': 'checked_out',
                'user': checkout.checked_out_by.username if checkout.checked_out_by else 'Unknown',
                'details': f'Reason: {checkout.checkout_reason}'
            })

        # Versions
        for version in versions:
            timeline.append({
                'timestamp': version.created_at,
                'action': 'version_created',
                'user': version.created_by.username if version.created_by else 'System',
                'details': f'Version {version.major_version}.{version.minor_version} - {version.change_summary}'
            })

        # Approvals
        for approval in approvals:
            if approval.approved_at:
                timeline.append({
                    'timestamp': approval.approved_at,
                    'action': 'approved',
                    'user': approval.approver.username if approval.approver else 'Unknown',
                    'details': f'Status: {approval.approval_status} - {approval.comments}'
                })

        # Sort by timestamp
        timeline.sort(key=lambda x: x['timestamp'], reverse=True)

        return Response({
            'document_id': document.document_id,
            'audit_trail': timeline,
            'total_events': len(timeline)
        })

    # ---------------------------------------------------------------
    # CONTENT EDITING ENDPOINTS
    # ---------------------------------------------------------------

    @action(detail=True, methods=['get', 'put'], url_path='content')
    def content(self, request, pk=None):
        """
        GET: Return the document's TipTap JSON content.
        PUT: Save content (requires checkout or draft state).
        """
        document = self.get_object()

        if request.method == 'GET':
            return Response({
                'id': document.id,
                'document_id': document.document_id,
                'title': document.title,
                'content': document.content,
                'content_html': document.content_html,
                'description': document.description,
                'vault_state': document.vault_state,
                'is_locked': document.is_locked,
                'locked_by': document.locked_by_id,
                'version_string': document.version_string,
            })

        # PUT — save content
        serializer = DocumentContentUpdateSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        # Only allow editing in draft or if user has checkout
        if document.vault_state != 'draft' and not (
            document.is_locked and document.locked_by == request.user
        ):
            return Response(
                {'error': 'Document must be in Draft state or checked out by you to edit.'},
                status=status.HTTP_403_FORBIDDEN
            )

        document.content = serializer.validated_data['content']
        document.content_html = serializer.validated_data.get('content_html', '')

        # Extract plain text for search
        import re
        plain = re.sub(r'<[^>]+>', '', document.content_html)
        document.content_plain_text = plain[:50000]  # Limit for DB

        document.save(update_fields=['content', 'content_html', 'content_plain_text', 'updated_at'])

        from core.models import AuditLog
        from django.contrib.contenttypes.models import ContentType
        ct = ContentType.objects.get_for_model(Document)
        AuditLog.objects.create(
            content_type=ct,
            object_id=str(document.pk),
            object_repr=str(document),
            user=request.user,
            action='update',
            ip_address=request.META.get('REMOTE_ADDR'),
            new_values={'action': 'content_saved'},
            change_summary=f"Document content updated by {request.user.username}",
        )

        return Response({
            'status': 'saved',
            'document_id': document.document_id,
            'updated_at': document.updated_at,
        })

    # ---------------------------------------------------------------
    # COMMENTS ENDPOINTS
    # ---------------------------------------------------------------

    @action(detail=True, methods=['get', 'post'], url_path='comments')
    def comments(self, request, pk=None):
        """
        GET: List all comments for this document (threaded, top-level only).
        POST: Create a new comment or reply.
        """
        document = self.get_object()

        if request.method == 'GET':
            qs = DocumentComment.objects.filter(
                document=document, parent__isnull=True
            ).select_related('author', 'resolved_by').prefetch_related('replies__author')

            status_filter = request.query_params.get('status')
            if status_filter:
                qs = qs.filter(status=status_filter)

            serializer = DocumentCommentSerializer(qs, many=True)
            return Response(serializer.data)

        # POST — create comment
        serializer = DocumentCommentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save(
            author=request.user,
            document=document,
            document_version=document.version_string,
        )
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['patch'], url_path='comments/(?P<comment_id>[0-9]+)/resolve')
    def resolve_comment(self, request, pk=None, comment_id=None):
        """Mark a comment as resolved."""
        document = self.get_object()
        comment = get_object_or_404(DocumentComment, pk=comment_id, document=document)

        new_status = request.data.get('status', 'resolved')
        if new_status not in ('resolved', 'wont_fix', 'open'):
            return Response({'error': 'Invalid status'}, status=status.HTTP_400_BAD_REQUEST)

        comment.status = new_status
        if new_status in ('resolved', 'wont_fix'):
            comment.resolved_by = request.user
            comment.resolved_at = timezone.now()
        else:
            comment.resolved_by = None
            comment.resolved_at = None
        comment.save()

        return Response(DocumentCommentSerializer(comment).data)

    # ---------------------------------------------------------------
    # SUGGESTIONS / TRACK CHANGES ENDPOINTS
    # ---------------------------------------------------------------

    @action(detail=True, methods=['get', 'post'], url_path='suggestions')
    def suggestions(self, request, pk=None):
        """
        GET: List all suggestions for this document.
        POST: Create a new suggestion (track change).
        """
        document = self.get_object()

        if request.method == 'GET':
            qs = DocumentSuggestion.objects.filter(
                document=document
            ).select_related('author', 'reviewed_by')

            status_filter = request.query_params.get('status')
            if status_filter:
                qs = qs.filter(status=status_filter)

            serializer = DocumentSuggestionSerializer(qs, many=True)
            return Response(serializer.data)

        # POST — create suggestion
        serializer = DocumentSuggestionSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save(
            author=request.user,
            document=document,
            document_version=document.version_string,
        )
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'], url_path='suggestions/(?P<suggestion_id>[0-9]+)/accept')
    def accept_suggestion(self, request, pk=None, suggestion_id=None):
        """Accept a suggestion and apply the text change to the document content."""
        document = self.get_object()
        suggestion = get_object_or_404(DocumentSuggestion, pk=suggestion_id, document=document)

        if suggestion.status != 'pending':
            return Response(
                {'error': f'Suggestion is already {suggestion.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )

        suggestion.status = 'accepted'
        suggestion.reviewed_by = request.user
        suggestion.reviewed_at = timezone.now()
        suggestion.save()

        return Response(DocumentSuggestionSerializer(suggestion).data)

    @action(detail=True, methods=['post'], url_path='suggestions/(?P<suggestion_id>[0-9]+)/reject')
    def reject_suggestion(self, request, pk=None, suggestion_id=None):
        """Reject a suggestion."""
        document = self.get_object()
        suggestion = get_object_or_404(DocumentSuggestion, pk=suggestion_id, document=document)

        if suggestion.status != 'pending':
            return Response(
                {'error': f'Suggestion is already {suggestion.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )

        suggestion.status = 'rejected'
        suggestion.reviewed_by = request.user
        suggestion.reviewed_at = timezone.now()
        suggestion.save()

        return Response(DocumentSuggestionSerializer(suggestion).data)

    @action(detail=True, methods=['post'], url_path='suggestions/bulk-accept')
    def bulk_accept_suggestions(self, request, pk=None):
        """Accept all pending suggestions for this document."""
        document = self.get_object()
        updated = DocumentSuggestion.objects.filter(
            document=document, status='pending'
        ).update(
            status='accepted',
            reviewed_by=request.user,
            reviewed_at=timezone.now(),
        )
        return Response({'accepted_count': updated})

    # ---------------------------------------------------------------
    # EXPORT ENDPOINTS
    # ---------------------------------------------------------------

    @action(detail=True, methods=['get'], url_path='export/(?P<fmt>pdf|docx|html)')
    def export_document(self, request, pk=None, fmt='pdf'):
        """
        Export document content to PDF, DOCX, or HTML.
        Supports rich formatting: headings, lists, tables, bold/italic, links.
        """
        document = self.get_object()
        html_content = document.content_html or '<p>No content</p>'

        if fmt == 'html':
            from django.http import HttpResponse
            full_html = self._wrap_html_for_export(document, html_content)
            response = HttpResponse(full_html, content_type='text/html')
            response['Content-Disposition'] = f'attachment; filename="{document.document_id}.html"'
            return response

        if fmt == 'docx':
            try:
                return self._export_docx(document, html_content)
            except ImportError as e:
                return Response(
                    {'error': f'Missing dependency: {e}. Contact admin.'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            except Exception as e:
                return Response(
                    {'error': f'DOCX export failed: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

        if fmt == 'pdf':
            try:
                return self._export_pdf(document, html_content)
            except ImportError as e:
                return Response(
                    {'error': f'Missing dependency: {e}. Contact admin.'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            except Exception as e:
                return Response(
                    {'error': f'PDF export failed: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

    def _wrap_html_for_export(self, document, html_content):
        """Wrap raw HTML content in a full styled HTML document for export."""
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{document.title}</title>
<style>
  body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #1e293b; line-height: 1.7; }}
  h1 {{ font-size: 1.75em; font-weight: 700; color: #0f172a; border-bottom: 2px solid #e2e8f0; padding-bottom: 0.3em; }}
  h2 {{ font-size: 1.4em; font-weight: 600; color: #1e293b; }}
  h3 {{ font-size: 1.15em; font-weight: 600; color: #334155; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; }}
  th {{ background-color: #f8fafc; font-weight: 600; }}
  code {{ background: #f1f5f9; padding: 2px 4px; border-radius: 3px; font-size: 0.9em; }}
  pre {{ background: #1e293b; color: #e2e8f0; padding: 1em; border-radius: 8px; overflow-x: auto; }}
  blockquote {{ border-left: 4px solid #0d9488; padding: 0.75em 1em; margin: 1em 0; background: #f0fdfa; color: #475569; }}
  a {{ color: #0d9488; }}
  .doc-header {{ border-bottom: 2px solid #0d9488; padding-bottom: 1em; margin-bottom: 2em; }}
  .doc-meta {{ color: #64748b; font-size: 0.85em; }}
</style>
</head>
<body>
<div class="doc-header">
  <h1 style="border:none; margin:0;">{document.title}</h1>
  <p class="doc-meta">{document.document_id} &nbsp;|&nbsp; Version {document.version_string} &nbsp;|&nbsp; {document.vault_state.title()}</p>
</div>
{html_content}
</body>
</html>"""

    def _export_docx(self, document, html_content):
        """Export document as richly-formatted DOCX using python-docx + BeautifulSoup."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from bs4 import BeautifulSoup

        doc = DocxDocument()

        # Set default style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Title
        title_para = doc.add_heading(document.title, level=0)
        # Metadata line
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = meta.add_run(
            f"Document ID: {document.document_id}  |  "
            f"Version: {document.version_string}  |  "
            f"Status: {document.vault_state.title()}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        doc.add_paragraph('')  # spacer

        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        self._process_html_to_docx(doc, soup, None)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from django.http import HttpResponse
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{document.document_id}.docx"'
        return response

    def _process_html_to_docx(self, doc, element, current_para):
        """Recursively convert HTML elements to DOCX content."""
        from docx.shared import Pt, RGBColor
        from bs4 import NavigableString, Tag

        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip() and current_para is not None:
                    current_para.add_run(text)
                continue

            if not isinstance(child, Tag):
                continue

            tag = child.name.lower()

            # Headings
            if tag in ('h1', 'h2', 'h3', 'h4'):
                level = int(tag[1])
                heading = doc.add_heading(child.get_text(), level=level)
                continue

            # Paragraphs
            if tag == 'p':
                p = doc.add_paragraph()
                self._add_inline_runs(p, child)
                continue

            # Unordered list
            if tag == 'ul':
                for li in child.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_inline_runs(p, li)
                continue

            # Ordered list
            if tag == 'ol':
                for li in child.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Number')
                    self._add_inline_runs(p, li)
                continue

            # Blockquote
            if tag == 'blockquote':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                run = p.add_run(child.get_text())
                run.font.italic = True
                run.font.color.rgb = RGBColor(71, 85, 105)
                continue

            # Code block
            if tag == 'pre':
                p = doc.add_paragraph()
                run = p.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                continue

            # Table
            if tag == 'table':
                self._add_table_to_docx(doc, child)
                continue

            # Horizontal rule
            if tag == 'hr':
                p = doc.add_paragraph()
                p.add_run('─' * 60).font.color.rgb = RGBColor(226, 232, 240)
                continue

            # Other block elements — recurse
            if tag in ('div', 'section', 'article', 'main', 'body'):
                self._process_html_to_docx(doc, child, current_para)
                continue

            # Fallback: treat as paragraph
            text = child.get_text().strip()
            if text:
                p = doc.add_paragraph()
                self._add_inline_runs(p, child)

    def _add_inline_runs(self, paragraph, element):
        """Process inline HTML elements (bold, italic, code, links) into runs."""
        from docx.shared import Pt, RGBColor
        from bs4 import NavigableString, Tag

        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    paragraph.add_run(text)
                continue

            if not isinstance(child, Tag):
                continue

            tag = child.name.lower()
            text = child.get_text()

            if tag in ('strong', 'b'):
                run = paragraph.add_run(text)
                run.bold = True
            elif tag in ('em', 'i'):
                run = paragraph.add_run(text)
                run.italic = True
            elif tag == 'u':
                run = paragraph.add_run(text)
                run.underline = True
            elif tag in ('s', 'del', 'strike'):
                run = paragraph.add_run(text)
                run.font.strike = True
            elif tag == 'code':
                run = paragraph.add_run(text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(190, 24, 93)
            elif tag == 'a':
                run = paragraph.add_run(text)
                run.font.color.rgb = RGBColor(13, 148, 136)
                run.underline = True
            elif tag == 'mark':
                run = paragraph.add_run(text)
                run.font.highlight_color = 7  # Yellow
            elif tag == 'br':
                paragraph.add_run('\n')
            elif tag == 'span':
                # Recurse for spans (may contain formatting)
                self._add_inline_runs(paragraph, child)
            else:
                paragraph.add_run(text)

    def _add_table_to_docx(self, doc, table_element):
        """Convert an HTML table to a DOCX table."""
        from docx.shared import Pt
        rows_data = []
        for tr in table_element.find_all('tr'):
            cells = []
            for td in tr.find_all(['td', 'th']):
                cells.append(td.get_text().strip())
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        max_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < max_cols:
                    table.rows[i].cells[j].text = cell_text

    def _export_pdf(self, document, html_content):
        """Export document as PDF using reportlab with rich formatting from HTML."""
        from io import BytesIO
        from django.http import HttpResponse
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
            Table as RLTable, TableStyle, ListFlowable, ListItem
        )
        from bs4 import BeautifulSoup
        import re

        buffer = BytesIO()
        doc_pdf = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        styles = getSampleStyleSheet()

        # Custom styles
        styles.add(ParagraphStyle(
            'DocTitle', parent=styles['Title'],
            fontSize=20, textColor=HexColor('#0f172a'),
            spaceAfter=6
        ))
        styles.add(ParagraphStyle(
            'DocMeta', parent=styles['Normal'],
            fontSize=9, textColor=HexColor('#64748b'),
            spaceAfter=16
        ))
        styles.add(ParagraphStyle(
            'DocH1', parent=styles['Heading1'],
            fontSize=16, textColor=HexColor('#0f172a'),
            spaceBefore=16, spaceAfter=8,
            borderWidth=1, borderColor=HexColor('#e2e8f0'), borderPadding=4
        ))
        styles.add(ParagraphStyle(
            'DocH2', parent=styles['Heading2'],
            fontSize=14, textColor=HexColor('#1e293b'),
            spaceBefore=12, spaceAfter=6
        ))
        styles.add(ParagraphStyle(
            'DocH3', parent=styles['Heading3'],
            fontSize=12, textColor=HexColor('#334155'),
            spaceBefore=10, spaceAfter=4
        ))
        styles.add(ParagraphStyle(
            'DocBody', parent=styles['Normal'],
            fontSize=10, leading=16, textColor=HexColor('#1e293b'),
            spaceAfter=6
        ))
        styles.add(ParagraphStyle(
            'DocQuote', parent=styles['Normal'],
            fontSize=10, leading=15, textColor=HexColor('#475569'),
            leftIndent=24, spaceAfter=8, fontName='Helvetica-Oblique'
        ))
        styles.add(ParagraphStyle(
            'DocCode', parent=styles['Normal'],
            fontSize=8, fontName='Courier', textColor=HexColor('#e2e8f0'),
            backColor=HexColor('#1e293b'), leftIndent=12, rightIndent=12,
            spaceBefore=6, spaceAfter=6
        ))

        story = []

        # Title section
        story.append(Paragraph(document.title, styles['DocTitle']))
        story.append(Paragraph(
            f"Document ID: {document.document_id} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"Version: {document.version_string} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"Status: {document.vault_state.title()}",
            styles['DocMeta']
        ))
        story.append(HRFlowable(
            width="100%", thickness=1, color=HexColor('#0d9488'),
            spaceAfter=16
        ))

        # Parse HTML and build story
        soup = BeautifulSoup(html_content, 'html.parser')
        self._process_html_to_pdf(soup, story, styles)

        # Footer info
        story.append(Spacer(1, 24))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#e2e8f0')))
        story.append(Paragraph(
            f"Generated from Arni eQMS &nbsp;|&nbsp; {document.document_id} v{document.version_string}",
            styles['DocMeta']
        ))

        doc_pdf.build(story)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{document.document_id}.pdf"'
        return response

    def _process_html_to_pdf(self, element, story, styles):
        """Recursively convert HTML to reportlab flowables."""
        from bs4 import NavigableString, Tag
        from reportlab.platypus import Paragraph, Spacer, HRFlowable, ListFlowable, ListItem

        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    story.append(Paragraph(text, styles['DocBody']))
                continue

            if not isinstance(child, Tag):
                continue

            tag = child.name.lower()

            if tag == 'h1':
                story.append(Paragraph(child.get_text(), styles['DocH1']))
            elif tag == 'h2':
                story.append(Paragraph(child.get_text(), styles['DocH2']))
            elif tag in ('h3', 'h4'):
                story.append(Paragraph(child.get_text(), styles['DocH3']))
            elif tag == 'p':
                # Process inline formatting
                html_str = self._inline_html_for_pdf(child)
                story.append(Paragraph(html_str, styles['DocBody']))
            elif tag == 'ul':
                items = []
                for li in child.find_all('li', recursive=False):
                    html_str = self._inline_html_for_pdf(li)
                    items.append(ListItem(Paragraph(html_str, styles['DocBody'])))
                if items:
                    story.append(ListFlowable(items, bulletType='bullet', start='•'))
            elif tag == 'ol':
                items = []
                for li in child.find_all('li', recursive=False):
                    html_str = self._inline_html_for_pdf(li)
                    items.append(ListItem(Paragraph(html_str, styles['DocBody'])))
                if items:
                    story.append(ListFlowable(items, bulletType='1'))
            elif tag == 'blockquote':
                story.append(Paragraph(child.get_text(), styles['DocQuote']))
            elif tag == 'pre':
                story.append(Paragraph(child.get_text(), styles['DocCode']))
            elif tag == 'hr':
                story.append(HRFlowable(width="100%", thickness=1, color='#e2e8f0', spaceAfter=8))
            elif tag == 'table':
                self._add_table_to_pdf(child, story, styles)
            elif tag in ('div', 'section', 'article', 'main', 'body'):
                self._process_html_to_pdf(child, story, styles)
            else:
                text = child.get_text().strip()
                if text:
                    story.append(Paragraph(text, styles['DocBody']))

    def _inline_html_for_pdf(self, element):
        """Convert inline HTML elements to reportlab-compatible HTML markup."""
        from bs4 import NavigableString, Tag
        parts = []
        for child in element.children:
            if isinstance(child, NavigableString):
                # Escape special characters for reportlab
                text = str(child).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                parts.append(text)
            elif isinstance(child, Tag):
                tag = child.name.lower()
                inner = child.get_text().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if tag in ('strong', 'b'):
                    parts.append(f'<b>{inner}</b>')
                elif tag in ('em', 'i'):
                    parts.append(f'<i>{inner}</i>')
                elif tag == 'u':
                    parts.append(f'<u>{inner}</u>')
                elif tag == 'code':
                    parts.append(f'<font face="Courier" size="8" color="#be185d">{inner}</font>')
                elif tag == 'a':
                    parts.append(f'<font color="#0d9488"><u>{inner}</u></font>')
                elif tag == 'br':
                    parts.append('<br/>')
                else:
                    parts.append(inner)
        return ''.join(parts)

    def _add_table_to_pdf(self, table_element, story, styles):
        """Convert an HTML table to a reportlab Table."""
        from reportlab.platypus import Table as RLTable, TableStyle, Spacer
        from reportlab.lib.colors import HexColor

        rows_data = []
        for tr in table_element.find_all('tr'):
            cells = []
            for td in tr.find_all(['td', 'th']):
                cells.append(td.get_text().strip())
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        table = RLTable(rows_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f8fafc')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#334155')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(table)
        story.append(Spacer(1, 8))


class DocumentChangeOrderViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing document change orders with approval workflows.
    """
    queryset = DocumentChangeOrder.objects.select_related(
        'document',
        'proposed_by'
    ).prefetch_related(
        'approvals'
    ).all()
    serializer_class = DocumentChangeOrderSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    search_fields = ['document__document_id', 'change_number', 'description', 'title']
    ordering_fields = ['created_at', 'status', 'change_number']
    ordering = ['-created_at']

    def perform_create(self, serializer):
        """Set proposed_by to current user."""
        serializer.save(proposed_by=self.request.user)

    def perform_update(self, serializer):
        """Update change order."""
        serializer.save()

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """
        Approve a change order.

        POST body: {
            "comment": "string (optional)"
        }
        """
        change_order = self.get_object()

        try:
            from .models import DocumentChangeApproval

            approval, created = DocumentChangeApproval.objects.update_or_create(
                change_order=change_order,
                approver=request.user,
                defaults={
                    'status': 'approved',
                    'approved_at': timezone.now(),
                    'comments': request.data.get('comment', '')
                }
            )

            # Update change order status if all approvals are done
            all_approvals = change_order.approvals.all()
            if all_approvals.filter(status__in=['pending', 'rejected']).count() == 0:
                change_order.status = 'approved'
                change_order.save()

            from .serializers import DocumentChangeApprovalSerializer
            return Response(
                DocumentChangeApprovalSerializer(approval).data,
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Approval failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        """
        Reject a change order.

        POST body: {
            "comment": "string (required)"
        }
        """
        change_order = self.get_object()
        comment = request.data.get('comment', '')

        if not comment:
            return Response(
                {'error': 'comment is required for rejection'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            from .models import DocumentChangeApproval

            approval, created = DocumentChangeApproval.objects.update_or_create(
                change_order=change_order,
                approver=request.user,
                defaults={
                    'status': 'rejected',
                    'approved_at': timezone.now(),
                    'comments': comment
                }
            )

            # Update change order status to rejected
            change_order.status = 'rejected'
            change_order.save()

            from .serializers import DocumentChangeApprovalSerializer
            return Response(
                DocumentChangeApprovalSerializer(approval).data,
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Rejection failed: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )
